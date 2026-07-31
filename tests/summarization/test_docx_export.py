from datetime import datetime, timezone

from docx import Document

from app.summarization.docx_export import build_docx_filename, export_mom_docx
from app.summarization.groq_client import MomResult


def test_build_docx_filename_replaces_spaces_with_dashes():
    filename = build_docx_filename(datetime(2026, 7, 31, 14, 0), "Test Meeting 3")
    assert filename == "2026-07-31-Test-Meeting-3.docx"


def test_build_docx_filename_strips_unsafe_characters():
    filename = build_docx_filename(datetime(2026, 7, 31), "Rapat: Q3 / Review?")
    assert " " not in filename
    assert all(c not in filename for c in ':/\\?*<>|"')


def test_build_docx_filename_uses_the_wib_date_not_the_utc_one():
    """Timestamps are stored UTC and displayed WIB. 2026-07-31 20:00 UTC is a
    2026-08-01 03:00 WIB meeting, so the file must be dated 2026-08-01."""
    filename = build_docx_filename(
        datetime(2026, 7, 31, 20, 0, tzinfo=timezone.utc), "Rapat Malam",
    )
    assert filename == "2026-08-01-Rapat-Malam.docx"


def test_docx_body_date_matches_the_filename_date(tmp_path):
    mom = MomResult(minute_by_minute=[], decisions=[], action_items=[], detailed_notes="-")
    meeting_date = datetime(2026, 7, 31, 20, 0, tzinfo=timezone.utc)
    output_path = tmp_path / build_docx_filename(meeting_date, "Rapat Malam")

    export_mom_docx("Rapat Malam", meeting_date, mom, output_path)

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "01 August 2026 03:00 WIB" in full_text
    assert output_path.name.startswith("2026-08-01")


def test_export_creates_docx_with_expected_sections(tmp_path):
    mom = MomResult(
        minute_by_minute=[{"time": "00:00", "point": "Pembukaan"}],
        decisions=["Lanjutkan rencana A"],
        action_items=[{"item": "Kirim laporan", "assignee": "Budi", "due": "2026-08-01"}],
        detailed_notes="Semua peserta setuju melanjutkan.",
    )
    output_path = tmp_path / "mom.docx"

    result_path = export_mom_docx("Rapat Mingguan", datetime(2026, 7, 30, 9, 0), mom, output_path)

    assert result_path == output_path
    assert output_path.exists()

    doc = Document(str(output_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Rapat Mingguan" in full_text
    assert "Ringkasan Menit ke Menit" in full_text
    assert "Keputusan" in full_text
    assert "Action Items" in full_text
    assert "Catatan Detail" in full_text
    assert "Lanjutkan rencana A" in full_text
    assert "Semua peserta setuju melanjutkan." in full_text
