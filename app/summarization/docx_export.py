import re
from datetime import datetime
from pathlib import Path

from docx import Document

from app.summarization.groq_client import MomResult


def build_docx_filename(meeting_date: datetime, meeting_title: str) -> str:
    """e.g. 2026-07-31-Test-Meeting-3.docx -- no spaces, safe on Windows."""
    slug = re.sub(r"\s+", "-", meeting_title.strip())
    slug = re.sub(r"[^A-Za-z0-9\-_]", "", slug) or "meeting"
    return f"{meeting_date.strftime('%Y-%m-%d')}-{slug}.docx"


def export_mom_docx(meeting_title: str, meeting_date: datetime, mom: MomResult, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading(meeting_title, level=0)
    doc.add_paragraph(f"Tanggal: {meeting_date.strftime('%d %B %Y %H:%M')}")

    doc.add_heading("Ringkasan Menit ke Menit", level=1)
    for entry in mom.minute_by_minute:
        doc.add_paragraph(f"{entry['time']} — {entry['point']}", style="List Bullet")

    doc.add_heading("Keputusan", level=1)
    for decision in mom.decisions:
        doc.add_paragraph(decision, style="List Bullet")

    doc.add_heading("Action Items", level=1)
    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text, header[1].text, header[2].text = "Item", "PIC", "Tenggat"
    for action in mom.action_items:
        row = table.add_row().cells
        row[0].text = action["item"]
        row[1].text = action["assignee"]
        row[2].text = action["due"]

    doc.add_heading("Catatan Detail", level=1)
    doc.add_paragraph(mom.detailed_notes)

    doc.save(str(output_path))
    return output_path
